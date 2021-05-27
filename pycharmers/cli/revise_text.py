#coding: utf-8
import os
import re
import sys
import argparse
import docx
import pandas as pd

from ..utils._colorings import toRED, toGREEN
from ..utils.generic_utils import now_str

def revise_text(argv=sys.argv[1:]):
    """Revise word file.

    Args:
        -W/--word (str)    : Path to the word file.
        -E/--excel (str)   : Path to the excel file.
        --sheet-name (str) : SheetName for the Excel.
        --NG (str)         : The column name which indicates NG words. Defaults to ``"NG"``
        --OK (str)         : The column name which indicates OK words. Defaults to ``"OK"``

    Note:
        When you run from the command line, execute as follows::
        
        $ revise-text --word sample.docx --excel wordlist.xlsx        
    """
    parser = argparse.ArgumentParser(prog="revise-text", add_help=True)
    parser.add_argument("-W", "--word",  type=str, help="Path to the word file.")
    parser.add_argument("-E", "--excel", type=str, help="Path to the excel file.")
    parser.add_argument("--sheet-name",  type=str, help="SheetName for the Excel.")
    parser.add_argument("--NG", type=str, default="NG", help="The column name which indicates NG words. Defaults to ``NG``")
    parser.add_argument("--OK", type=str, default="OK", help="The column name which indicates OK words. Defaults to ``OK``")
    args = parser.parse_args(argv)

    df_wordlist = pd.read_excel(args.excel, sheet_name=args.sheet_name).fillna("")
    NG2OK = dict(df_wordlist[[args.NG, args.OK]].values)

    def repl_create(para_no:int, para_digit:int, text_digit:int) -> callable:
        """Replace word while printing the logs.

        Args:
            para_no (int)    : The paragraph number.
            para_digit (int) : The digit for the number of paragraphs.
            text_digit (int) : The maximum digit for the number of texts in a paragraph.

        Returns:
            callable: Replacement function for ``re.sub``
        """
        def repl(match):
            ng_word = match.group()
            ok_word = NG2OK[ng_word]
            print(f"\t[Para.{para_no:>0{para_digit}}] ({match.start():>0{text_digit}}-{match.end():>0{text_digit}}) Reveised {toRED(ng_word)} -> {toGREEN(ok_word)}")
            return ok_word
        return repl

    doc = docx.Document(args.word)
    para_digit = len(str(len(doc.paragraphs)))
    text_digit = len(str(max([len(t) for t in doc.paragraphs])))
    for i,para in enumerate(doc.paragraphs):
        for ng in NG2OK.keys():
            para.text = re.sub(pattern=ng, repl=repl_create(para_no=i, para_digit=para_digit, text_digit=text_digit), string=para.text)
    doc.save(now_str().join(os.path.splitext(args.word)))
